#!/usr/bin/env python3
"""Align RIMS_Project_Report.docx with the implemented codebase."""

from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.table import Table

SRC = Path("/Users/ananyasrivastava/Downloads/RIMS_Project_Report.docx")
OUT = Path(
    "/Users/ananyasrivastava/Development/refugee-identity-system/RIMS_Project_Report_Updated.docx"
)


def replace_in_cell(cell, replacements: list[tuple[str, str]]) -> None:
    for p in cell.paragraphs:
        for old, new in replacements:
            if old in p.text:
                p.text = p.text.replace(old, new)


def replace_all(doc: Document, replacements: list[tuple[str, str]]) -> None:
    for p in doc.paragraphs:
        for old, new in replacements:
            if old in p.text:
                p.text = p.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, replacements)


def set_table_row(table: Table, row_idx: int, values: list[str]) -> None:
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        if i < len(row.cells):
            row.cells[i].text = val


def delete_table_row(table: Table, row_idx: int) -> None:
    tbl = table._tbl
    tr = table.rows[row_idx]._tr
    tbl.remove(tr)


def main() -> None:
    doc = Document(str(SRC))

    global_replacements = [
        ("PyTeal / PuyaPy", "PuyaPy (Algorand Python)"),
        ("PyTeal/PuyaPy", "PuyaPy"),
        ("PuyaPy / PyTeal", "PuyaPy (Algorand Python)"),
        (
            "production-grade, decentralized identity platform",
            "prototype decentralized identity platform",
        ),
        (
            "achieving full GDPR and eIDAS compliance",
            "following privacy-by-design principles aligned with GDPR and eIDAS goals (demo stores refugee profiles in local JSON files off-chain)",
        ),
        (
            "Stateless orchestration: biometric processing, SDK coordination",
            "Stateless orchestration: Algorand SDK coordination, migration workflow APIs, custodial wallet provisioning (liveness runs client-side in the browser)",
        ),
        (
            "biometric liveness, Algorand SDK coordinator, permission manager",
            "Algorand SDK coordinator, custodial wallet provisioning, migration and registry APIs",
        ),
        (
            "At any subsequent verification point — a border, a hospital, an aid distribution center — the verifying party scans the QR code, the backend reads the on-chain commitment, and the refugee's identity status is confirmed in under four seconds, with no internet connection to any centralized NGO database required.",
            "At verification time, an authorized party uses the FastAPI backend (with Algorand algod/indexer access) to read on-chain local state for the refugee wallet. The printed QR encodes identity_id and custodial W1 for migration handoff; aid distribution uses the on-chain claim_aid flag. Full offline border verification is a roadmap item — the current demo requires network access to the backend and Algorand nodes.",
        ),
        (
            "The RefugeeContract is deployed as ARC-4 compliant bytecode on Algorand TestNet. It exposes five callable functions with strict access control:",
            "The RefugeeContract is compiled with PuyaPy to ARC-4 bytecode and deployed via AlgoKit / the Admin portal (LocalNet or TestNet; App ID persisted in .deployments.json). It exposes four ABI methods used by the application:",
        ),
        (
            "impossible because H_person is generated server-side only upon successful completion of all five stages.",
            "blocked in practice because registration commits personhood_hash only after the aid-worker registration UI completes all five client-side liveness stages.",
        ),
        (
            "RIMS implements attribute-level consent management. A medical clinic requesting age verification receives only the binary result of H_age comparison — over18:true or over18:false — without access to the refugee's actual date of birth, name, or any other field. Every disclosure request is logged on-chain with requesting organization identifier, requested attributes, and refugee approval status, creating an immutable consent audit trail. Future deployments will integrate zk-SNARKs to enable zero-knowledge age proofs without even revealing the comparison result.",
            "The codebase includes a data-access request model stored in local JSON (.access-requests.json) and an Aid Worker “Request Access” page; the refugee consent UI exists but is not yet wired to a route. On-chain selective disclosure and zk-SNARK age proofs are documented as future scope — only commitment hashes and aid/migration state are anchored on-chain today.",
        ),
        (
            "RIMS achieves GDPR compliance through privacy-by-design: no PII on-chain, right to erasure for off-chain data, and explicit consent for every disclosure. eIDAS compliance is achieved through cryptographic identity binding equivalent to qualified electronic signatures.",
            "RIMS follows privacy-by-design: no raw PII on-chain (only SHA-256 commitment hashes in local state). Off-chain demographic fields live in .registry.json and can be removed in a deployment that implements erasure workflows. Portal authentication is demo-grade (localStorage / ID-only refugee login); production GDPR/eIDAS certification would require hardened auth, encryption at rest for custodial keys, and completed consent flows.",
        ),
        (
            "The FastAPI backend is containerized via Docker and orchestrated by Kubernetes, with each container running a single FastAPI worker capable of handling 100 concurrent requests. SQLite is used in development; PostgreSQL with high-availability replication is used in production. The full system is deployable on Algorand Mainnet without code changes — only environment configuration differs.",
            "The FastAPI backend runs as a local development service (npm run api on port 8000). Workflow state (registry, custodial wallets, migration requests, audit aggregates) is persisted in JSON files under Nexathon/, not a production database. Docker/AlgoKit LocalNet is used for contract integration tests. The same contract and backend can target TestNet or Mainnet by configuring DEPLOYER_MNEMONIC and algod endpoints.",
        ),
        (
            "The RefugeeContract is compiled from Python source using the Puya compiler and deployed as ARC-4 compliant bytecode on Algorand TestNet with App ID 758845823.",
            "The RefugeeContract is compiled from PuyaPy using AlgoKit and deployed to LocalNet or TestNet; the active App ID is written to Nexathon/.deployments.json (reference TestNet deployment: 758845823).",
        ),
        (
            "verified against a live Algorand TestNet deployment:",
            "verified against Algorand LocalNet via pytest (algokit localnet start):",
        ),
        (
            "All critical contract functions verified against live Algorand TestNet deployment.",
            "All critical contract functions verified on LocalNet integration tests.",
        ),
        (
            "establish RIMS as a production-ready architecture, not a theoretical proof of concept",
            "establish RIMS as a working end-to-end prototype with a clear path to production hardening",
        ),
        (
            "The empirical results — 98.2% spoofing rejection, 3.8-second transaction finality, $0.0004 per operation, 100% double-claim prevention under concurrent load, and 8/8 contract test cases passing",
            "The implemented results — client-side liveness gating at registration, Algorand finality on the order of seconds, low per-transaction fees on Algorand, duplicate aid-claim prevention in contract tests, and 8/8 LocalNet contract test cases passing",
        ),
        (
            "The system is open-source and designed for pilot deployment in humanitarian field operations, with ongoing empirical validation planned across East Africa and Southeast Asia.",
            "The system is designed for hackathon/demo deployment and further hardening before humanitarian field pilots.",
        ),
        (
            "The Refugee Portal gives refugees direct visibility into and control over their identity records. The landing dashboard displays their on-chain identity status, commitment hash summaries, aid claim history, and pending data-access requests. Refugees without a smartphone initiate migration by submitting a request from the portal (no wallet connect required); the Aid Worker later completes Pera Wallet signing in Wallet Migration Tools. A medical clinic requesting age verification sees only over18:true or over18:false without receiving the actual date of birth. The Wallet Migration Wizard guides refugees through a three-screen flow: QR scan of the physical RIMS Security Card to extract W1, Pera Wallet connection for W2, and the cryptographic challenge-response signing flow. Post-submission, a migration status tracker polls the backend every ten seconds to display progress through Pending Admin Review, Approved, and On-Chain Complete states.",
            "The Refugee Portal (/refugee) supports ID-only login for custodial identities (backend lookup in .custodial-wallets.json). Dashboard, identity details, and blockchain status pages load data via POST /api/blockchain/get-identity. Refugees submit wallet migration intent via POST /api/blockchain/migration-request-lite without connecting a wallet. Pera Wallet connection and challenge signing are performed by the aid worker in Wallet Migration Tools (/aid-worker/migration). Migration status is shown from backend JSON state; on-chain migrate_wallet runs when an admin approves and W2 has opted into the app.",
        ),
        (
            "The Admin Portal provides governance-level oversight of the entire RIMS deployment. The main dashboard aggregates real-time system metrics — total registrations, active identities, migrations completed, pending migration requests, aid distributions, and double-claim attempts — sourced directly from the blockchain and refreshed every 30 seconds. The Registrar Management widget displays all authorized aid workers with registration counts and last activity timestamps; administrators can revoke registrar permissions with a single action that immediately propagates to the RefugeeContract's Box Storage. The Pending Migrations Queue presents all awaiting wallet migration requests with refugee ID, proposed W2 address, submission timestamp, and the cryptographic signature verification result — enabling informed single-click approval or rejection. The Audit Trail Visualization renders an immutable event log of every on-chain transaction with clickable Algorand Explorer links.",
            "The Admin Portal (/admin) provides deploy controls, registrar authorization (add_registrar on-chain), aggregated stats (GET /api/admin/stats), refugee list, audit log derived from local JSON files, and migration approve/reject (POST /api/blockchain/migration-approve). Metrics combine registry/custodial JSON with optional indexer reads — not a live 30-second blockchain poll. Migration approval triggers migrate_wallet when W2 is opted in; otherwise the request may be marked approved in JSON only.",
        ),
        (
            "The Aid Worker Portal is engineered for field deployment in resource-constrained humanitarian environments. Its centerpiece is the Registration workflow: capture demographic fields, run the 5-stage MediaPipe liveness check in-browser, choose wallet type (Pera self-custody or “No smartphone” custodial W1), provision the wallet via backend APIs, persist the profile to the registry, and print a QR-encoded RIMS Security Card. Additional tools include refugee search, aid distribution (single on-chain aid_claimed flag), migration request inbox, and Wallet Migration Tools for Pera signing.",
            "The Aid Worker Portal (/aid-worker) implements registration with client-side MediaPipe liveness, custodial W1 provisioning (POST /api/blockchain/generate-custodial-wallet) or Pera-assisted register, QR card generation, refugee search, aid distribution via claim_aid, read-only migration requests list, and Wallet Migration Tools for challenge signing. Demo login uses localStorage (including admin/1234 backdoor) — not production authentication.",
        ),
        (
            "5 — Cryptographic Handshake\nRefugee W2\nPera Wallet signs challenge with W2 private key using signData",
            "5 — Cryptographic Handshake\nAid Worker (Pera W2)\nAid worker connects Pera Wallet and signs the migration challenge for the refugee’s new W2 address",
        ),
        (
            "3 — QR Scan\nRefugee\nScans physical card; extracts identity_id and W1 address",
            "3 — Migration intent\nRefugee\nSubmits migration-request-lite with identity_id (custodial lookup); optional QR encodes identity_id + W1 for aid-worker tools",
        ),
        (
            "The EAR threshold of 0.20 was empirically calibrated against 1,000 video sequences (500 genuine, 500 spoofing attempts), achieving optimal separation between live subjects and presentation attacks.",
            "The EAR threshold of 0.20 is used in the Register.jsx liveness UI (blink and head-yaw challenges). Formal large-scale spoofing benchmarks are planned; the pipeline is implemented client-side and does not send raw video to the server.",
        ),
        (
            "EMPIRICAL RESULT\nEAR threshold 0.20 rejects 98.2% of spoofing attacks (2D photos, silicone masks, replay video).\nMulti-stage sequential challenges defeat pre-recorded video attacks through temporal\ndiscontinuity detection across directed head pose transitions.",
            "IMPLEMENTATION NOTE\nLiveness runs entirely in the browser during aid-worker registration. The backend exposes POST /api/refugee/liveness-hash as a lightweight acknowledgment; personhood_hash on-chain is derived from the client liveness composite hash when valid.",
        ),
        (
            "The stress test result is particularly significant: under 100 concurrent aid claims within a 10-minute window, every transaction achieved finality with zero double-claims detected. This validates the smart contract's atomic aid-claimed flag mechanism as a reliable fraud prevention layer independent of backend database integrity.",
            "Contract integration tests on LocalNet confirm that a second claim_aid call for the same refugee wallet is rejected after aid_claimed is set — demonstrating on-chain duplicate-aid prevention independent of JSON registry state.",
        ),
        (
            "Zero (hashes only; off-chain PII shredded)",
            "Low (hashes on-chain; profiles in local JSON off-chain)",
        ),
        (
            "Attribute-level; on-chain audit trail",
            "Planned (JSON access requests; consent UI not routed)",
        ),
        (
            "Full (right to erasure; no PII on-chain)",
            "Partial (no PII on-chain; demo auth/storage)",
        ),
        (
            "Yes (QR card + cached ledger state)",
            "Partial (QR card; requires API + algod for verification)",
        ),
    ]
    replace_all(doc, global_replacements)

    # --- Architecture table: 3 layers (implemented stack) ---
    arch = doc.tables[6]
    # Keep header row; replace rows 1-4 with 3 layers
    set_table_row(
        arch,
        1,
        [
            "Layer 1 — Frontend",
            "React 18 + Vite + Tailwind: Aid Worker (/aid-worker), Refugee (/refugee), Admin (/admin); Pera Wallet via @perawallet/connect",
            "Browser client; demo localStorage auth",
        ],
    )
    set_table_row(
        arch,
        2,
        [
            "Layer 2 — Backend",
            "FastAPI (main.py): deploy/register/claim-aid/migrate_wallet orchestration, custodial W1 keys, JSON registry & migration stores",
            "Trusted orchestrator; signs transactions with DEPLOYER_MNEMONIC",
        ],
    )
    set_table_row(
        arch,
        3,
        [
            "Layer 3 — Blockchain",
            "Algorand AVM; RefugeeContract (PuyaPy); per-account local state + BoxMap registrars",
            "Cryptographic anchor for hashes, aid_claimed, migration",
        ],
    )
    delete_table_row(arch, 5)
    delete_table_row(arch, 4)

    # Update figure caption for architecture
    replace_in_cell(
        doc.tables[5].rows[0].cells[0],
        [
            (
                "Figure 1 — RIMS Five-Layer System Architecture Diagram",
                "Figure 1 — RIMS Three-Layer System Architecture Diagram",
            ),
            (
                "(Paste architecture diagram from the paper here)",
                "(Paste architecture diagram: React portals → FastAPI → Algorand RefugeeContract)",
            ),
        ],
    )

    # Fix section heading in paragraphs
    replace_all(
        doc,
        [
            ("4.1  Five-Layer Architecture", "4.1  Three-Layer Architecture"),
            (
                "RIMS is structured as a five-layer system with strict trust boundaries between each layer:",
                "RIMS is structured as a three-layer system matching the implemented repository:",
            ),
        ],
    )

    # --- Tech stack table ---
    stack = doc.tables[8]
    stack_rows = [
        ["Layer", "Technology", "Version", "Purpose"],
        [
            "Smart Contract",
            "PuyaPy (Algorand Python)",
            "AVM 10",
            "RefugeeContract: register, claim_aid, migrate_wallet, add_registrar",
        ],
        [
            "Contract Framework",
            "AlgoKit",
            "Latest",
            "Build, LocalNet/TestNet deploy, typed Python client generation",
        ],
        [
            "Blockchain",
            "Algorand LocalNet / TestNet",
            "AVM",
            "On-chain local state + registrar BoxMap",
        ],
        [
            "Wallet",
            "Pera Wallet + @perawallet/connect",
            "Latest",
            "Aid-worker migration signing; optional refugee self-custody at registration",
        ],
        [
            "Frontend",
            "React 18 + TypeScript + Vite + TailwindCSS",
            "Latest",
            "Aid Worker, Refugee, and Admin portals",
        ],
        [
            "Biometric",
            "MediaPipe Face Mesh (browser)",
            "Latest",
            "Client-side liveness during registration (no video sent to server)",
        ],
        [
            "Backend",
            "FastAPI (Python)",
            "Latest",
            "Blockchain orchestration, custodial W1, JSON-backed workflow APIs",
        ],
        [
            "Persistence",
            "JSON files (.registry.json, .custodial-wallets.json, .migration-requests.json, .deployments.json)",
            "Demo",
            "Registry, custodial keys, migration queue (no PostgreSQL in repo)",
        ],
        [
            "Security (roadmap)",
            "Encrypt custodial keys; portal JWT/OAuth",
            "Planned",
            "Demo stores W1 private_key_b64 in .custodial-wallets.json",
        ],
    ]
    while len(stack.rows) > len(stack_rows):
        delete_table_row(stack, len(stack.rows) - 1)
    for i, values in enumerate(stack_rows):
        set_table_row(stack, i, values)

    # --- Remove log_permission row from contract table ---
    contract_table = doc.tables[11]
    delete_table_row(contract_table, 5)

    # --- Performance table: keep structure, qualify unbenchmarked rows ---
    perf = doc.tables[18]
    note_rows = {1, 2, 3, 4, 5, 6, 7, 10}  # biometric + stress rows (0-indexed data rows)
    for idx in note_rows:
        if idx < len(perf.rows):
            sig_cell = perf.rows[idx].cells[3]
            if "Design target" not in sig_cell.text and "LocalNet" not in sig_cell.text:
                sig_cell.text = (
                    sig_cell.text.split("\n")[0]
                    + " (design target — not formally benchmarked in this codebase)"
                )
    set_table_row(
        perf,
        10,
        [
            "Contract tests",
            "LocalNet pytest suite (test_contract.py)",
            "8/8 passed",
            "Duplicate register, duplicate aid claim, and migration rejected as expected",
        ],
    )

    # --- Insert note before section 9 if paragraph exists ---
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("9.  Empirical Performance Results"):
            p.text = (
                "9.  Performance & Validation Results\n"
                "Metrics below mix Algorand network characteristics, implemented client-side liveness behavior, "
                "and design targets. Biometric accuracy percentages are illustrative until a formal evaluation dataset is run."
            )
            break

    # --- Future scope: mark zk-SNARK consent as future (already is) ---
    replace_all(
        doc,
        [
            (
                "Integrate zk-SNARKs/zk-STARKs for attribute proofs without revealing values",
                "Integrate zk-SNARKs/zk-STARKs for attribute proofs (consent UI and on-chain disclosure not in current build)",
            ),
        ],
    )

    # --- Direct cell fixes (text split across runs broke global replace) ---
    doc.tables[16].rows[5].cells[1].text = "Aid Worker (Pera W2)"
    doc.tables[16].rows[5].cells[2].text = (
        "Aid worker connects Pera Wallet and signs the migration challenge "
        "for the refugee’s proposed W2 address (signData)"
    )
    doc.tables[17].rows[0].cells[0].text = "IMPLEMENTATION NOTE"
    doc.tables[17].rows[1].cells[0].text = (
        "Liveness runs in the browser during aid-worker registration. "
        "POST /api/refugee/liveness-hash acknowledges completion; "
        "personhood_hash on-chain uses the client composite when valid."
    )
    doc.tables[18].rows[6].cells[2].text = "Target (not benchmarked)"
    doc.tables[18].rows[6].cells[3].text = (
        "Illustrative spoofing benchmark — formal evaluation pending"
    )

    for p in doc.paragraphs:
        if "1,000 video sequences" in p.text:
            p.text = (
                "The EAR threshold of 0.20 is enforced in the Register.jsx liveness UI "
                "(blink and head-yaw challenges). Large-scale spoofing benchmarks are planned."
            )
        if p.text.startswith("The Aid Worker Portal is engineered"):
            p.text = (
                "The Aid Worker Portal (/aid-worker) supports refugee registration with a "
                "client-side MediaPipe liveness flow in Register.jsx, custodial W1 provisioning "
                "via POST /api/blockchain/generate-custodial-wallet (or Pera-assisted register), "
                "QR security card generation, refugee search, aid distribution (on-chain claim_aid), "
                "read-only migration requests, and Wallet Migration Tools where the aid worker "
                "connects Pera to sign the migration challenge. Authentication is demo-grade "
                "(localStorage passwords; admin/1234 backdoor)."
            )
        if p.text.startswith("The Refugee Portal gives"):
            p.text = (
                "The Refugee Portal (/refugee) lets custodial refugees log in with identity_id only "
                "(POST /api/blockchain/verify-identity). Pages include dashboard, identity details, "
                "blockchain status, and migration request (POST /api/blockchain/migration-request-lite) "
                "without wallet connect. Pera signing happens in aid-worker tools; admin approval "
                "completes on-chain migrate_wallet when W2 has opted in."
            )
        if p.text.startswith("The Admin Portal provides"):
            p.text = (
                "The Admin Portal (/admin) includes dashboard stats (GET /api/admin/stats), audit log "
                "(GET /api/audit/logs from JSON aggregates), refugee list, contract deploy and "
                "add_registrar (AdminStatus.jsx), and migration approve/reject. There is no "
                "production admin login in the demo build."
            )

    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
